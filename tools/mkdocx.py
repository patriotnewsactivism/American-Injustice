# -*- coding: utf-8 -*-
import io, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION

SRC = "/home/user/American-Injustice/manuscript-exports/American_Injustice_FULL.md"
OUT = "/home/user/American-Injustice/manuscript-exports/American_Injustice_FULL_13pt.docx"
BODY_PT = 13.0          # 11 -> 12 -> 13pt
SERIF = "Georgia"   # heavier than EB Garamond; Don found the lighter face too thin

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(6.0), Inches(9.0)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin = sec.bottom_margin = Inches(0.8)

st = doc.styles["Normal"]
st.font.name = SERIF
st.font.size = Pt(BODY_PT)
pf = st.paragraph_format
pf.space_after = Pt(0)
pf.first_line_indent = Inches(0.22)
pf.line_spacing = 1.18

for name, size, before, after in (("Heading 1", 20, 0, 14),
                                  ("Heading 2", 14, 16, 8),
                                  ("Heading 3", 12.5, 12, 6)):
    s = doc.styles[name]
    s.font.name = SERIF
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.first_line_indent = Inches(0)
    s.paragraph_format.keep_with_next = True

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.S)

def emit_runs(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(BODY_PT - 1.5)
        else:
            p.add_run(tok)

import sys, re as _re; sys.path.insert(0,"/tmp/claude-0/-home-user/bb9b3891-5cc9-5d86-9295-b7ae9ed35788/scratchpad")
from prep import load
_md, PART_AT = load()
PART_AT[9999]="PART VI \u00b7 SYSTEMIC ANALYSIS"
lines = _md.split("\n")

# ---- title page ----
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(150); t.paragraph_format.first_line_indent = Inches(0)
r = t.add_run("AMERICAN INJUSTICE"); r.bold = True; r.font.size = Pt(30); r.font.name = SERIF
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.paragraph_format.space_before = Pt(18); s.paragraph_format.first_line_indent = Inches(0)
r = s.add_run("Don Matthews"); r.font.size = Pt(14); r.font.name = SERIF
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

for _i, (_txt, _sc, _sp) in enumerate([
        (u"For Lydia-Elise, Anna-Claire, McKenna-Rose,\nEthan-Oliver, and Meri-Emmelyn \u2014\nwho never deserved any of this.", False, 190),
        (u"For Andy Arant \u2014\nattorney, friend, and mentor,\nlost to cancer.", False, 22),
        (u"And for Bradley Foust,\nwhere it all started.", False, 22),
        (u"REST IN PEACE.  SEMPER FI.", True, 30)]):
    _p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p.paragraph_format.first_line_indent = Inches(0)
    _p.paragraph_format.space_before = Pt(_sp)
    for _k, _line in enumerate(_txt.split("\n")):
        if _k: _p.add_run().add_break()
        _r = _p.add_run(_line); _r.font.size = Pt(11 if _sc else 12); _r.font.name = SERIF
        if _sc: _r.font.small_caps = True
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

i, first_h1, tbl = 0, True, None
while i < len(lines):
    ln = lines[i].rstrip()

    # markdown table
    if ln.startswith("|") and i + 1 < len(lines) and set(lines[i+1].replace("|","").strip()) <= set("-: "):
        hdr = [c.strip() for c in ln.strip("|").split("|")]
        rows = []
        i += 2
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([c.strip() for c in lines[i].strip("|").split("|")])
            i += 1
        tb = doc.add_table(rows=1, cols=len(hdr)); tb.style = "Table Grid"
        for k, h in enumerate(hdr):
            cp = tb.rows[0].cells[k].paragraphs[0]; cp.paragraph_format.first_line_indent = Inches(0)
            rr = cp.add_run(re.sub(r"\*", "", h)); rr.bold = True; rr.font.size = Pt(BODY_PT - 1.5); rr.font.name = SERIF
        for row in rows:
            cells = tb.add_row().cells
            for k, c in enumerate(row[:len(hdr)]):
                cp = cells[k].paragraphs[0]; cp.paragraph_format.first_line_indent = Inches(0)
                emit_runs(cp, c)
                for rn in cp.runs:
                    rn.font.size = Pt(BODY_PT - 1.5); rn.font.name = SERIF
        doc.add_paragraph()
        continue

    if not ln.strip():
        i += 1; continue

    if ln.startswith("# "):
        _plain = re.sub(r"\*", "", ln[2:].strip())
        _cm = _re.match(r"CHAPTER\s+(\d+)", _plain, _re.I)
        _pn = int(_cm.group(1)) if _cm else None
        _part = PART_AT.get(_pn) or (PART_AT[9999] if _plain.upper().startswith("LEGAL ANALYSIS") else None)
        if not first_h1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        first_h1 = False
        if _part:
            _pp = doc.add_paragraph(); _pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pp.paragraph_format.first_line_indent = Inches(0); _pp.paragraph_format.space_before = Pt(210)
            _r = _pp.add_run(_part); _r.bold = True; _r.font.size = Pt(15); _r.font.name = SERIF
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph(style="Heading 1"); emit_runs(p, ln[2:].strip())
    elif ln.startswith("## "):
        p = doc.add_paragraph(style="Heading 2"); emit_runs(p, ln[3:].strip())
    elif ln.startswith("### "):
        p = doc.add_paragraph(style="Heading 3"); emit_runs(p, ln[4:].strip())
    elif ln.strip() in ("---", "***", "___"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
        p.add_run("⁂")
    elif ln.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(0); emit_runs(p, ln[2:])
        for rn in p.runs: rn.italic = True
    elif re.match(r"^[-*] ", ln) or re.match(r"^\d+\. ", ln):
        p = doc.add_paragraph(style="List Bullet" if ln[0] in "-*" else "List Number")
        p.paragraph_format.first_line_indent = Inches(0)
        emit_runs(p, re.sub(r"^([-*]|\d+\.) ", "", ln))
        for rn in p.runs: rn.font.size = Pt(BODY_PT); rn.font.name = SERIF
    else:
        p = doc.add_paragraph()
        # section-note italic blocks and standalone bold headers get no indent
        if ln.startswith("*") and ln.endswith("*") and not ln.startswith("**"):
            p.paragraph_format.first_line_indent = Inches(0)
        if ln.startswith("**") and ln.endswith("**"):
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(12)
        emit_runs(p, ln)
    i += 1

doc.save(OUT)
print("wrote", OUT)
